"""
File Layer — text and structure extraction for PDF and DOCX.
Returns ExtractedDocument regardless of input format.
"""

from pathlib import Path
from typing import List

from backend.models.schemas import ExtractedDocument, FileType


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def extract_pdf(path: Path) -> ExtractedDocument:
    """
    Extract text from a PDF.

    Raises:
        FileNotFoundError: if the file does not exist.
        ValueError: if the file extension is not .pdf.
    """
    import fitz  # PyMuPDF

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if path.suffix.lower() != ".pdf":
        raise ValueError(f"Expected .pdf extension, got: {path.suffix}")

    doc = fitz.open(str(path))

    # Password-protected detection
    if doc.needs_pass:
        doc.close()
        return ExtractedDocument(
            filename=path.name,
            file_type=FileType.PDF,
            body_text="",
            header_text="",
            footer_text="",
            page_count=0,
            is_scanned=False,
            is_password_protected=True,
        )

    pages_text: List[str] = []
    for page in doc:
        pages_text.append(page.get_text("text"))

    doc.close()

    body_text = "\n".join(pages_text).strip()
    is_scanned = len(body_text) == 0

    return ExtractedDocument(
        filename=path.name,
        file_type=FileType.PDF,
        body_text=body_text,
        header_text="",
        footer_text="",
        page_count=len(pages_text),
        is_scanned=is_scanned,
        is_password_protected=False,
    )


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_docx(path: Path) -> ExtractedDocument:
    """
    Extract text from a DOCX file, including headers, footers, and table cells.

    Raises:
        FileNotFoundError: if the file does not exist.
    """
    from docx import Document as DocxDocument

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    doc = DocxDocument(str(path))

    # Body paragraphs (exclude table paragraphs — those come via table_cells)
    body_parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            body_parts.append(para.text)

    # Table cells — preserved as 2-D list
    table_cells: List[List[str]] = []
    for table in doc.tables:
        for row in table.rows:
            table_cells.append([cell.text for cell in row.cells])

    # Headers and footers (first section)
    header_text = ""
    footer_text = ""
    if doc.sections:
        section = doc.sections[0]
        header_paras = section.header.paragraphs if section.header else []
        footer_paras = section.footer.paragraphs if section.footer else []
        header_text = "\n".join(p.text for p in header_paras if p.text.strip())
        footer_text = "\n".join(p.text for p in footer_paras if p.text.strip())

    # Page count approximation (python-docx has no native page count)
    page_count = max(1, len(body_parts) // 10 + 1)

    body_text = "\n".join(body_parts)

    return ExtractedDocument(
        filename=path.name,
        file_type=FileType.DOCX,
        body_text=body_text,
        header_text=header_text,
        footer_text=footer_text,
        page_count=page_count,
        is_scanned=False,
        is_password_protected=False,
        table_cells=table_cells if table_cells else None,
    )


# ---------------------------------------------------------------------------
# Generic dispatcher
# ---------------------------------------------------------------------------

def extract_file(path: Path) -> ExtractedDocument:
    """
    Dispatch to the correct extractor based on file extension.

    Raises:
        ValueError: for unsupported file types.
        FileNotFoundError: if the file does not exist.
    """
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    elif ext == ".docx":
        return extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
