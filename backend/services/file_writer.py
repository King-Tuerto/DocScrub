"""
Format-preserving output writers for PDF and DOCX.

PDF  — PyMuPDF redaction: find PII text → add redact annotation with placeholder
       text → apply_redactions() removes original and overlays the replacement.
DOCX — python-docx paragraph replacement: rebuild each paragraph's runs so the
       full text reflects the replacement while preserving paragraph-level style
       (alignment, spacing, indent) and table structure.

Both anonymise and re-identify paths are supported by swapping the lookup direction.
"""

from pathlib import Path
from typing import Dict


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def write_anonymized_pdf(input_path: Path, output_path: Path, mapping_table) -> None:
    """
    Produce an anonymized PDF from *input_path* by finding each PII string on every
    page and replacing it with the corresponding placeholder via redaction.
    """
    lookup = _build_lookup(mapping_table, reverse=False)
    _process_pdf(input_path, output_path, lookup)


def write_reidentified_pdf(input_path: Path, output_path: Path, mapping_table) -> None:
    """Replace placeholder tokens with original values in a PDF."""
    lookup = _build_lookup(mapping_table, reverse=True)
    _process_pdf(input_path, output_path, lookup)


def _build_lookup(mapping_table, reverse: bool) -> Dict[str, str]:
    """Build {find: replace} from a MappingTable, sorted longest-first."""
    if reverse:
        pairs = [(e.placeholder, e.original) for e in mapping_table.entries
                 if e.placeholder and e.original]
    else:
        pairs = [(e.original, e.placeholder) for e in mapping_table.entries
                 if e.original and e.placeholder]
    # Longest-match-first so "Jane Smith" wins over "Jane"
    pairs.sort(key=lambda x: -len(x[0]))
    return dict(pairs)


def _process_pdf(input_path: Path, output_path: Path, lookup: Dict[str, str]) -> None:
    import fitz  # PyMuPDF

    doc = fitz.open(str(input_path))

    if not lookup:
        doc.save(str(output_path))
        doc.close()
        return

    for page in doc:
        for find_text, replace_text in lookup.items():
            if not find_text:
                continue
            areas = page.search_for(find_text)
            for area in areas:
                page.add_redact_annot(
                    area,
                    text=replace_text,
                    fontsize=11,
                    fill=(1, 1, 1),       # white background
                    text_color=(0, 0, 0), # black text
                )
        page.apply_redactions()

    doc.save(str(output_path))
    doc.close()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def write_anonymized_docx(input_path: Path, output_path: Path, mapping_table) -> None:
    """Produce an anonymized DOCX with PII replaced by placeholders."""
    lookup = _build_lookup(mapping_table, reverse=False)
    _process_docx(input_path, output_path, lookup)


def write_reidentified_docx(input_path: Path, output_path: Path, mapping_table) -> None:
    """Replace placeholder tokens with original values in a DOCX."""
    lookup = _build_lookup(mapping_table, reverse=True)
    _process_docx(input_path, output_path, lookup)


def _process_docx(input_path: Path, output_path: Path, lookup: Dict[str, str]) -> None:
    from docx import Document

    doc = Document(str(input_path))

    if lookup:
        # Sort pairs longest-first for correct replacement order
        sorted_pairs = sorted(lookup.items(), key=lambda x: -len(x[0]))
        _replace_in_docx(doc, sorted_pairs)

    doc.save(str(output_path))


def _replace_in_docx(doc, sorted_pairs) -> None:
    """
    Walk every paragraph in the document (body, tables, headers, footers) and
    apply text replacements.  Preserves paragraph-level style; per-run formatting
    is simplified to the first run's style.
    """
    def replace_para(para):
        full_text = para.text
        if not full_text:
            return
        new_text = full_text
        for orig, repl in sorted_pairs:
            if orig in new_text:
                new_text = new_text.replace(orig, repl)
        if new_text == full_text:
            return
        # Rebuild: put everything in the first run, clear the rest
        runs = para.runs
        if runs:
            runs[0].text = new_text
            for run in runs[1:]:
                run.text = ""
        else:
            para.add_run(new_text)

    # Body paragraphs
    for para in doc.paragraphs:
        replace_para(para)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_para(para)

    # Headers and footers (all sections)
    for section in doc.sections:
        try:
            if section.header:
                for para in section.header.paragraphs:
                    replace_para(para)
        except Exception:
            pass
        try:
            if section.footer:
                for para in section.footer.paragraphs:
                    replace_para(para)
        except Exception:
            pass
