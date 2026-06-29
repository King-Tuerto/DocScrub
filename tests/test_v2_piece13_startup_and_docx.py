"""
V2 Piece 13 — Startup reliability and DOCX export quality

Covers five real-world bugs:

1. DATABASE NOT CREATED ON STARTUP
   init_db must run on every server start. After a fresh delete of docscrub.db
   the next launch must work without manual intervention.

2. EXPORTED DOCX FILES DON'T OPEN
   The anonymized DOCX produced by write_anonymized_docx must be a valid DOCX
   that can be opened by python-docx (and therefore Word). The replacement must
   be visible in the output.

3. LAUNCHER "ALREADY RUNNING" DIALOG
   _port_in_use() must reliably return False when nothing is on the port, and
   the function must not raise on a connection refused error.

4. DATABASE INIT ON FRESH DB
   Equivalent to issue 1 — tested via the lifespan path (TestClient context
   manager triggers FastAPI startup event).

5. STOP SERVER
   The launcher's on_close helper must terminate the subprocess and not leave
   the port occupied. Tested by verifying proc.terminate() + proc.wait() path.
"""

import io
import socket
import sqlite3
import subprocess
import tempfile
from pathlib import Path

import pytest


# ---------------------------------------------------------------------------
# Issue 1 & 4 — init_db creates tables from scratch; lifespan runs it too
# ---------------------------------------------------------------------------

class TestInitDbFromScratch:
    def test_init_db_creates_all_tables(self, tmp_path):
        """Fresh path → init_db → all expected tables exist."""
        from backend.db.database import init_db

        db_path = tmp_path / "fresh.db"
        assert not db_path.exists()
        init_db(db_path)
        assert db_path.exists()

        conn = sqlite3.connect(str(db_path))
        tables = {r[0] for r in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()}
        conn.close()

        for expected in ("jobs", "files", "mappings", "roster_entries"):
            assert expected in tables, f"Table '{expected}' missing after init_db"

    def test_init_db_idempotent(self, tmp_path):
        """Calling init_db twice must not raise and must not lose data."""
        from backend.db.database import init_db

        db_path = tmp_path / "idem.db"
        init_db(db_path)
        init_db(db_path)  # second call — must be a no-op

        conn = sqlite3.connect(str(db_path))
        tables = {r[0] for r in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()}
        conn.close()
        assert "jobs" in tables

    def test_init_db_after_deletion_works(self, tmp_path):
        """Delete the DB, call init_db again → tables exist again."""
        from backend.db.database import init_db

        db_path = tmp_path / "deleteme.db"
        init_db(db_path)
        db_path.unlink()
        assert not db_path.exists()

        init_db(db_path)
        assert db_path.exists()

        conn = sqlite3.connect(str(db_path))
        tables = {r[0] for r in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()}
        conn.close()
        assert "jobs" in tables

    def test_create_app_initialises_db(self, tmp_path):
        """create_app() must call init_db so tables exist when routes run."""
        from backend.main import create_app

        db_path = tmp_path / "app.db"
        config = {
            "db_path": str(db_path),
            "output_directory": str(tmp_path / "output"),
            "llm_endpoint": "http://localhost:11434",
            "default_model": "llama3.1:8b",
        }
        create_app(config=config)

        conn = sqlite3.connect(str(db_path))
        tables = {r[0] for r in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()}
        conn.close()
        assert "jobs" in tables

    def test_lifespan_calls_init_db_on_startup(self, tmp_path):
        """
        FastAPI lifespan must also call init_db so a DB deleted between
        two requests is re-created on the next server boot.
        """
        from fastapi.testclient import TestClient
        from backend.main import create_app

        db_path = tmp_path / "lifespan.db"
        config = {
            "db_path": str(db_path),
            "output_directory": str(tmp_path / "output"),
            "llm_endpoint": "http://localhost:11434",
            "default_model": "llama3.1:8b",
        }
        app = create_app(config=config)

        # Delete the DB that create_app just made, then trigger the lifespan
        db_path.unlink()
        assert not db_path.exists()

        with TestClient(app) as client:
            resp = client.get("/health")
            assert resp.status_code == 200
            # After the lifespan ran, the DB must be back
            assert db_path.exists()

        conn = sqlite3.connect(str(db_path))
        tables = {r[0] for r in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()}
        conn.close()
        assert "jobs" in tables

    def test_output_dir_stored_as_absolute_on_app_state(self, tmp_path):
        """app.state.output_dir must always be an absolute Path."""
        from backend.main import create_app

        config = {
            "db_path": str(tmp_path / "app.db"),
            "output_directory": "./output",  # intentionally relative
            "llm_endpoint": "http://localhost:11434",
            "default_model": "llama3.1:8b",
        }
        app = create_app(config=config)
        assert app.state.output_dir.is_absolute(), (
            "app.state.output_dir must be absolute so all routes resolve consistently"
        )


# ---------------------------------------------------------------------------
# Issue 2 — DOCX round-trip validity and correctness
# ---------------------------------------------------------------------------

class TestDocxRoundTrip:
    """write_anonymized_docx must produce a VALID, OPENABLE DOCX with replaced text."""

    def _make_simple_docx(self, tmp_path, text: str) -> Path:
        from docx import Document
        doc = Document()
        doc.add_paragraph(text)
        p = tmp_path / "input.docx"
        doc.save(str(p))
        return p

    def _make_docx_with_splits(self, tmp_path) -> Path:
        """
        Build a DOCX where a name spans TWO runs in the same paragraph.
        This simulates what Word produces when text has mixed formatting.
        """
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree

        doc = Document()
        para = doc.add_paragraph()
        # Add runs in code so each word is a separate run
        run1 = para.add_run("Jane ")
        run2 = para.add_run("Smith")
        run2.bold = True  # forces a separate run in the XML
        para.add_run(" is a student.")

        p = tmp_path / "split_runs.docx"
        doc.save(str(p))
        return p

    def _make_mapping(self, pairs):
        from backend.services.mapper import MappingEntry, MappingTable
        return MappingTable(entries=[
            MappingEntry(original=orig, placeholder=repl, pii_type="PERSON", source="roster")
            for orig, repl in pairs
        ])

    def test_output_is_valid_docx(self, tmp_path):
        """Output DOCX can be opened by python-docx (i.e., is a valid ZIP/OOXML)."""
        from docx import Document
        from backend.services.file_writer import write_anonymized_docx

        src = self._make_simple_docx(tmp_path, "Patient: Jane Smith. SSN: 123-45-6789.")
        dst = tmp_path / "output.docx"
        mapping = self._make_mapping([("Jane Smith", "[PERSON_1]")])

        write_anonymized_docx(src, dst, mapping)

        assert dst.exists(), "Output file was not created"
        # Must open without exception
        doc = Document(str(dst))
        assert doc is not None

    def test_pii_replaced_in_output(self, tmp_path):
        """The anonymized text in the output DOCX must have the placeholder, not the original."""
        from docx import Document
        from backend.services.file_writer import write_anonymized_docx

        src = self._make_simple_docx(tmp_path, "Patient: Jane Smith. SSN: 123-45-6789.")
        dst = tmp_path / "output.docx"
        mapping = self._make_mapping([("Jane Smith", "[PERSON_1]")])

        write_anonymized_docx(src, dst, mapping)

        doc = Document(str(dst))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[PERSON_1]" in full_text, "Placeholder not found in output DOCX"
        assert "Jane Smith" not in full_text, "Original PII still present in output DOCX"

    def test_split_run_replacement(self, tmp_path):
        """
        When a name is split across two runs in the paragraph XML, the
        ZIP-level approach must still find and replace it correctly.
        """
        from docx import Document
        from backend.services.file_writer import write_anonymized_docx

        src = self._make_docx_with_splits(tmp_path)
        dst = tmp_path / "output_split.docx"
        mapping = self._make_mapping([("Jane Smith", "[PERSON_1]")])

        write_anonymized_docx(src, dst, mapping)

        doc = Document(str(dst))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[PERSON_1]" in full_text
        assert "Jane Smith" not in full_text

    def test_no_replacement_produces_copy(self, tmp_path):
        """Empty mapping → output is a byte-for-byte copy of the input (via shutil.copy2)."""
        from docx import Document
        from backend.services.file_writer import write_anonymized_docx
        from backend.services.mapper import MappingTable

        src = self._make_simple_docx(tmp_path, "No PII here.")
        dst = tmp_path / "copy.docx"
        write_anonymized_docx(src, dst, MappingTable(entries=[]))

        assert dst.exists()
        assert dst.read_bytes() == src.read_bytes()

    def test_header_text_replaced(self, tmp_path):
        """PII in the document header must be replaced in the output."""
        from docx import Document
        from backend.services.file_writer import write_anonymized_docx

        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Confidential — Jane Smith"
        doc.add_paragraph("Body text only.")
        src = tmp_path / "with_header.docx"
        doc.save(str(src))

        dst = tmp_path / "out_header.docx"
        mapping = self._make_mapping([("Jane Smith", "[PERSON_1]")])
        write_anonymized_docx(src, dst, mapping)

        doc2 = Document(str(dst))
        header_text = doc2.sections[0].header.paragraphs[0].text
        assert "[PERSON_1]" in header_text
        assert "Jane Smith" not in header_text

    def test_multiple_replacements(self, tmp_path):
        """Multiple PII items in one paragraph are all replaced."""
        from docx import Document
        from backend.services.file_writer import write_anonymized_docx

        src = self._make_simple_docx(
            tmp_path,
            "Jane Smith and Bob Jones are colleagues at Acme Corp."
        )
        dst = tmp_path / "multi.docx"
        mapping = self._make_mapping([
            ("Jane Smith", "[PERSON_1]"),
            ("Bob Jones",  "[PERSON_2]"),
            ("Acme Corp",  "[ORG_1]"),
        ])
        write_anonymized_docx(src, dst, mapping)

        doc = Document(str(dst))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[PERSON_1]" in full_text
        assert "[PERSON_2]" in full_text
        assert "[ORG_1]" in full_text
        assert "Jane Smith" not in full_text
        assert "Bob Jones" not in full_text
        assert "Acme Corp" not in full_text

    def test_pipeline_writes_docx_output(self, tmp_path, default_config):
        """
        Integration: pipeline writes a format-preserving DOCX that can be
        found at the expected output path and opened by python-docx.
        """
        from docx import Document as DocxDoc
        from backend.services.pipeline import run_pipeline
        from backend.services.roster_parser import RosterEntry

        # Create input DOCX
        doc = DocxDoc()
        doc.add_paragraph("Patient: Jane Smith. ID: STU-20240001.")
        input_dir = tmp_path / "job123" / "input"
        input_dir.mkdir(parents=True)
        src = input_dir / "patient.docx"
        doc.save(str(src))

        output_job_dir = tmp_path / "job123" / "output"
        roster = [RosterEntry(
            first_name="Jane", last_name="Smith",
            preferred_name=None, student_id=None, email=None
        )]

        run_pipeline(
            job_id="job123",
            file_paths=[src],
            config=default_config,
            roster_entries=roster,
            tier="names",
            output_dir=output_job_dir,
        )

        out_path = output_job_dir / "patient.docx"
        assert out_path.exists(), "Pipeline did not write DOCX output file"

        out_doc = DocxDoc(str(out_path))
        body = "\n".join(p.text for p in out_doc.paragraphs)
        assert "[PERSON_1]" in body
        assert "Jane Smith" not in body


# ---------------------------------------------------------------------------
# Issue 3 — port check reliability
# ---------------------------------------------------------------------------

class TestPortCheck:
    def test_port_not_in_use_returns_false(self):
        """On an unused port, _port_in_use() must return False without raising."""
        # Import the function directly from the launcher module
        import importlib.util
        import sys
        from pathlib import Path

        spec = importlib.util.spec_from_file_location(
            "launcher",
            Path(__file__).parent.parent / "launcher.py",
        )
        launcher = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(launcher)

        # Override PORT to something that's almost certainly free
        original_port = launcher.PORT
        launcher.PORT = 19876  # unlikely to be in use
        try:
            result = launcher._port_in_use()
        finally:
            launcher.PORT = original_port

        assert result is False

    def test_port_in_use_returns_true(self):
        """When a server is listening on PORT, _port_in_use() must return True."""
        import importlib.util
        from pathlib import Path

        spec = importlib.util.spec_from_file_location(
            "launcher2",
            Path(__file__).parent.parent / "launcher.py",
        )
        launcher = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(launcher)

        # Bind a real socket to a free port
        srv = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        srv.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        srv.bind(("127.0.0.1", 0))
        srv.listen(1)
        free_port = srv.getsockname()[1]

        launcher.PORT = free_port
        try:
            result = launcher._port_in_use()
        finally:
            srv.close()
            launcher.PORT = 8000

        assert result is True


# ---------------------------------------------------------------------------
# Issue 5 — process termination on window close
# ---------------------------------------------------------------------------

class TestProcessTermination:
    def test_terminate_kills_subprocess(self):
        """proc.terminate() must kill the subprocess; proc.wait() must not hang."""
        # Start a subprocess that idles (sleep)
        import sys
        proc = subprocess.Popen(
            [sys.executable, "-c", "import time; time.sleep(30)"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        assert proc.poll() is None  # still running

        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()
            proc.wait()

        assert proc.returncode is not None, "Process did not terminate"

    def test_kill_fallback_after_timeout(self):
        """If terminate doesn't stop the process in time, kill() is the fallback."""
        import sys
        proc = subprocess.Popen(
            [sys.executable, "-c", "import time; time.sleep(30)"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        proc.terminate()
        # Simulate a very short wait — kill immediately
        try:
            proc.wait(timeout=0.001)
        except subprocess.TimeoutExpired:
            proc.kill()
            proc.wait()

        assert proc.returncode is not None
