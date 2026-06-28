"""
Shared fixtures for all DocScrub tests.

Creates test artifacts (PDFs, DOCX files, DB paths, mock HTTP responses)
using third-party libraries that are project dependencies.  Production
modules are intentionally NOT imported here — each test file imports
what it needs so failures are isolated.
"""

import io
import json
import os
import sqlite3
import sys
from pathlib import Path

import pytest

# ---------------------------------------------------------------------------
# Path setup — make `backend` importable from every test
# ---------------------------------------------------------------------------

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


# ---------------------------------------------------------------------------
# Known-PII constants used across multiple test files
# ---------------------------------------------------------------------------

KNOWN_PII = {
    "person":   "Jane Smith",
    "org":      "Acme Corporation",
    "email":    "jane.smith@acme.com",
    "phone1":   "(555) 123-4567",
    "phone2":   "555-987-6543",
    "phone3":   "555.222.3333",
    "phone4":   "+15551234567",
    "ssn":      "123-45-6789",
    "address":  "742 Evergreen Terrace, Springfield, IL 62701",
    "zip":      "62701",
    "account":  "ACCT-9988-7766",
    "student_id": "STU-20240001",
    "dob":      "Date of Birth: 03/14/1985",
}

SAMPLE_PII_TEXT = f"""
CONFIDENTIAL DOCUMENT

Patient: {KNOWN_PII['person']}
Employer: {KNOWN_PII['org']}
Email: {KNOWN_PII['email']}
Phone: {KNOWN_PII['phone1']}
Mobile: {KNOWN_PII['phone2']}
SSN: {KNOWN_PII['ssn']}
Address: {KNOWN_PII['address']}
Account: {KNOWN_PII['account']}
Student ID: {KNOWN_PII['student_id']}
{KNOWN_PII['dob']}

Please contact {KNOWN_PII['person']} at {KNOWN_PII['email']} or {KNOWN_PII['phone1']}.
This document belongs to {KNOWN_PII['org']}.
""".strip()

CLEAN_TEXT = """
This document discusses general policy matters.
All quarterly targets were met.
The committee voted unanimously to approve the resolution.
No further action is required at this time.
""".strip()

VALID_LLM_JSON = json.dumps([
    {"text": KNOWN_PII["person"],  "type": "PERSON",  "confidence": "high"},
    {"text": KNOWN_PII["org"],     "type": "ORG",     "confidence": "high"},
    {"text": KNOWN_PII["email"],   "type": "EMAIL",   "confidence": "high"},
    {"text": KNOWN_PII["phone1"],  "type": "PHONE",   "confidence": "high"},
    {"text": KNOWN_PII["address"], "type": "ADDRESS", "confidence": "high"},
    {"text": KNOWN_PII["account"], "type": "ACCOUNT", "confidence": "medium"},
    {"text": KNOWN_PII["dob"],     "type": "DOB",     "confidence": "medium"},
    {"text": KNOWN_PII["student_id"], "type": "ID",   "confidence": "high"},
])

GARBAGE_LLM_RESPONSE = "Sure! Here is a summary of the document contents..."

EMPTY_LLM_RESPONSE = "[]"

PARTIAL_VALID_LLM_RESPONSE = json.dumps([
    {"text": KNOWN_PII["person"], "type": "PERSON"},     # missing 'confidence'
    {"text": KNOWN_PII["email"]},                         # missing 'type'
])

OVERLAPPING_PII_TEXT = (
    f"Contact {KNOWN_PII['person']} at {KNOWN_PII['email']} — "
    f"both name and email appear in the same sentence."
)

TABLE_PII_TEXT = """
| Name         | SSN           | Email                  |
|--------------|---------------|------------------------|
| Jane Smith   | 123-45-6789   | jane.smith@acme.com    |
| Bob Jones    | 987-65-4321   | bob.jones@example.com  |
"""


# ---------------------------------------------------------------------------
# Fixtures: text
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_pii_text():
    return SAMPLE_PII_TEXT


@pytest.fixture
def clean_text():
    return CLEAN_TEXT


@pytest.fixture
def overlapping_pii_text():
    return OVERLAPPING_PII_TEXT


@pytest.fixture
def table_pii_text():
    return TABLE_PII_TEXT


@pytest.fixture
def large_text():
    """Text long enough to require chunking (>4096 tokens ~ 16 000 chars)."""
    paragraph = SAMPLE_PII_TEXT + "\n\n"
    # repeat until we have plenty of content
    return paragraph * 80


# ---------------------------------------------------------------------------
# Fixtures: mock LLM API responses
# ---------------------------------------------------------------------------

@pytest.fixture
def valid_llm_response():
    return VALID_LLM_JSON


@pytest.fixture
def garbage_llm_response():
    return GARBAGE_LLM_RESPONSE


@pytest.fixture
def empty_llm_response():
    return EMPTY_LLM_RESPONSE


@pytest.fixture
def partial_valid_llm_response():
    return PARTIAL_VALID_LLM_RESPONSE


# ---------------------------------------------------------------------------
# Fixtures: file paths
# ---------------------------------------------------------------------------

@pytest.fixture
def tmp_db_path(tmp_path):
    return tmp_path / "test_docscrub.db"


@pytest.fixture
def output_dir(tmp_path):
    d = tmp_path / "output"
    d.mkdir()
    return d


@pytest.fixture
def mappings_dir(tmp_path):
    d = tmp_path / "mappings"
    d.mkdir()
    return d


# ---------------------------------------------------------------------------
# Fixtures: DOCX factory
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_docx_path(tmp_path):
    """Creates a .docx with known PII in body, header, footer, and a table."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()

    # Header
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"HEADER: {KNOWN_PII['org']} — Confidential"

    # Footer
    footer = section.footer
    footer.paragraphs[0].text = (
        f"FOOTER: {KNOWN_PII['person']} | {KNOWN_PII['phone1']}"
    )

    # Body paragraphs
    doc.add_paragraph(f"Patient Name: {KNOWN_PII['person']}")
    doc.add_paragraph(f"Email: {KNOWN_PII['email']}")
    doc.add_paragraph(f"SSN: {KNOWN_PII['ssn']}")
    doc.add_paragraph(f"Address: {KNOWN_PII['address']}")
    doc.add_paragraph(f"Account: {KNOWN_PII['account']}")
    doc.add_paragraph(f"Student ID: {KNOWN_PII['student_id']}")
    doc.add_paragraph(KNOWN_PII["dob"])

    # Table with PII
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "SSN"
    table.cell(0, 2).text = "Email"
    table.cell(1, 0).text = KNOWN_PII["person"]
    table.cell(1, 1).text = KNOWN_PII["ssn"]
    table.cell(1, 2).text = KNOWN_PII["email"]

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def multi_page_docx_path(tmp_path):
    """DOCX with enough content to test large-doc chunking."""
    from docx import Document

    doc = Document()
    for i in range(50):
        doc.add_paragraph(
            f"Page section {i}: {KNOWN_PII['person']} | "
            f"{KNOWN_PII['email']} | SSN: {KNOWN_PII['ssn']}"
        )
    path = tmp_path / "large.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Fixtures: PDF factory (PyMuPDF)
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_pdf_path(tmp_path):
    """Creates a text-layer PDF with known PII using PyMuPDF."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (50, 50),
        SAMPLE_PII_TEXT,
        fontsize=11,
    )
    path = tmp_path / "sample.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def scanned_pdf_path(tmp_path):
    """PDF containing only an embedded image — no text layer (simulates a scan)."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    # Insert a tiny solid rectangle as a stand-in image; no text inserted
    page.draw_rect(fitz.Rect(50, 50, 200, 200), color=(0, 0, 0), fill=(0.8, 0.8, 0.8))
    path = tmp_path / "scanned.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def password_protected_pdf_path(tmp_path):
    """PDF encrypted with a user password."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), SAMPLE_PII_TEXT, fontsize=11)
    path = tmp_path / "protected.pdf"
    doc.save(str(path), encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="secret")
    doc.close()
    return path


@pytest.fixture
def pdf_with_image_path(tmp_path):
    """PDF with an embedded raster image (for image extraction tests)."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Document with logo below.", fontsize=11)

    # Create a tiny 10x10 red PNG in memory
    img_doc = fitz.open()
    img_page = img_doc.new_page(width=10, height=10)
    img_page.draw_rect(fitz.Rect(0, 0, 10, 10), color=(1, 0, 0), fill=(1, 0, 0))
    img_bytes = img_doc.convert_to_pdf()
    img_doc.close()

    # Embed as an image reference (simplest approach: embed another PDF page as xobject)
    src = fitz.open("pdf", img_bytes)
    page.show_pdf_page(fitz.Rect(100, 100, 200, 200), src, 0)
    src.close()

    path = tmp_path / "with_image.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def docx_with_image_path(tmp_path):
    """DOCX with an embedded PNG image."""
    import struct
    import zlib
    from docx import Document

    # Minimal valid 1x1 red PNG (hand-crafted bytes)
    def make_minimal_png():
        def chunk(name, data):
            c = name + data
            crc = zlib.crc32(c) & 0xFFFFFFFF
            return struct.pack(">I", len(data)) + c + struct.pack(">I", crc)

        png = b"\x89PNG\r\n\x1a\n"
        ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        png += chunk(b"IHDR", ihdr_data)
        raw = b"\x00\xff\x00\x00"  # filter byte + RGB pixel
        compressed = zlib.compress(raw)
        png += chunk(b"IDAT", compressed)
        png += chunk(b"IEND", b"")
        return png

    doc = Document()
    doc.add_paragraph("Document with embedded image.")
    img_stream = io.BytesIO(make_minimal_png())
    doc.add_picture(img_stream)
    path = tmp_path / "with_image.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Fixtures: config
# ---------------------------------------------------------------------------

@pytest.fixture
def default_config():
    return {
        "llm_endpoint": "http://localhost:11434",
        "default_model": "llama3.1:8b",
        "output_directory": "./output",
        "db_path": "./docscrub.db",
        "image_review_default": "remove",
        "custom_regex_patterns": [],
    }


@pytest.fixture
def config_file(tmp_path, default_config):
    path = tmp_path / "config.json"
    path.write_text(json.dumps(default_config))
    return path


# ---------------------------------------------------------------------------
# Fixtures: FastAPI test client
# ---------------------------------------------------------------------------

@pytest.fixture
def app_client(tmp_path, default_config):
    """Returns a FastAPI TestClient backed by a temp DB and output dir."""
    from fastapi.testclient import TestClient
    from backend.main import create_app

    config = dict(default_config)
    config["db_path"] = str(tmp_path / "test.db")
    config["output_directory"] = str(tmp_path / "output")

    app = create_app(config=config)
    return TestClient(app)


# ---------------------------------------------------------------------------
# Fixtures: httpx mock helpers (used in LLM client tests)
# ---------------------------------------------------------------------------

@pytest.fixture
def mock_llm_endpoint(httpx_mock, valid_llm_response):
    """
    Registers a mock for the Ollama chat completions endpoint.
    Requires pytest-httpx to be installed.
    """
    httpx_mock.add_response(
        method="POST",
        url="http://localhost:11434/v1/chat/completions",
        json={
            "choices": [
                {"message": {"content": valid_llm_response}}
            ]
        },
    )
    httpx_mock.add_response(
        method="GET",
        url="http://localhost:11434/api/tags",
        json={
            "models": [
                {"name": "llama3.1:8b"},
                {"name": "mistral:7b"},
            ]
        },
    )
    return httpx_mock


@pytest.fixture
def mock_llm_garbage(httpx_mock, garbage_llm_response):
    httpx_mock.add_response(
        method="POST",
        url="http://localhost:11434/v1/chat/completions",
        json={"choices": [{"message": {"content": garbage_llm_response}}]},
    )
    return httpx_mock


@pytest.fixture
def mock_llm_unreachable(httpx_mock):
    import httpx
    httpx_mock.add_exception(
        httpx.ConnectError("Connection refused"),
        url="http://localhost:11434/v1/chat/completions",
    )
    return httpx_mock
