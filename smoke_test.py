#!/usr/bin/env python3
"""
DocScrub end-to-end smoke test.

Polls /health until the server is ready, then exercises the full request
cycle. Exit code 0 = all tests passed, 1 = at least one failed.

Usage: python smoke_test.py          (called by smoke_test.bat)
       python smoke_test.py --no-shutdown  (keep server alive after tests)
"""

import io
import json
import sys
import time

import fitz  # PyMuPDF
import httpx
from docx import Document

BASE = "http://127.0.0.1:8000"
_SHUTDOWN = "--no-shutdown" not in sys.argv

# PII embedded in the DOCX test
FIRST1, LAST1 = "Valentina", "Rosenstein"
FIRST2, LAST2 = "Marcus", "Okonkwo"
DOCX_EMAIL = "valentina.rosenstein@smoketest.example"
DOCX_PHONE = "555-867-5309"

# PII embedded in the PDF test (reuses roster names so Tier 2 catches them)
PDF_EMAIL = "pdf.test@smoketest.example"
PDF_PHONE = "555-234-5678"

_results: list[tuple[str, bool, str]] = []


def _ok(label: str) -> None:
    print(f"  [PASS] {label}")
    _results.append((label, True, ""))


def _fail(label: str, detail: str) -> None:
    print(f"  [FAIL] {label}")
    print(f"         {detail}")
    _results.append((label, False, detail))


def run(label: str, fn):
    """Run fn(); record PASS/FAIL.  Returns the fn() return value or None on failure."""
    try:
        result = fn()
        _ok(label)
        return result
    except Exception as exc:
        _fail(label, str(exc))
        return None


# ---------------------------------------------------------------------------
# Health-check polling
# ---------------------------------------------------------------------------

def _wait_for_server(timeout: int = 30) -> None:
    """Poll GET /health until the server responds or timeout expires."""
    print("  Waiting for server...", flush=True)
    deadline = time.monotonic() + timeout
    last_exc = None
    while time.monotonic() < deadline:
        try:
            httpx.get(f"{BASE}/health", timeout=2).raise_for_status()
            print("  Server ready.", flush=True)
            return
        except Exception as exc:
            last_exc = exc
            time.sleep(1)
    print(f"  ERROR: server did not respond within {timeout}s. Last error: {last_exc}")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph(
        f"Meeting notes for {FIRST1} {LAST1} and {FIRST2} {LAST2}."
    )
    doc.add_paragraph(f"Primary contact: {DOCX_EMAIL}")
    doc.add_paragraph(f"Call back at: {DOCX_PHONE}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_full_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def _make_pdf_bytes() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), f"Report by {FIRST1} {LAST1} and {FIRST2} {LAST2}.")
    page.insert_text((72, 130), f"Contact: {PDF_EMAIL}")
    page.insert_text((72, 160), f"Phone: {PDF_PHONE}")
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def _pdf_full_text(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    doc.close()
    return text


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    _wait_for_server()
    client = httpx.Client(base_url=BASE, timeout=120.0)

    roster_id: str | None = None
    job_id: str | None = None

    # ------------------------------------------------------------------
    # A. Create roster
    # ------------------------------------------------------------------
    def step_a():
        nonlocal roster_id
        resp = client.post("/rosters", json={"name": "Smoke Test Roster"})
        assert resp.status_code == 200, f"POST /rosters {resp.status_code}: {resp.text}"
        roster_id = resp.json()["id"]

        csv_bytes = (
            "first_name,last_name\n"
            f"{FIRST1},{LAST1}\n"
            f"{FIRST2},{LAST2}\n"
        ).encode()
        resp2 = client.post(
            f"/rosters/{roster_id}/entries",
            files={"file": ("roster.csv", csv_bytes, "text/csv")},
        )
        assert resp2.status_code == 200, (
            f"POST /rosters/.../entries {resp2.status_code}: {resp2.text}"
        )
        count = resp2.json()["count"]
        assert count == 2, f"Expected 2 entries, got {count}"

    run("A. Create roster with 2 names", step_a)

    if roster_id is None:
        print("\n  Cannot continue without a roster. Aborting.")
        _summary()
        client.close()
        return

    # ------------------------------------------------------------------
    # B. Upload test DOCX
    # ------------------------------------------------------------------
    def step_b():
        nonlocal job_id
        docx_bytes = _make_docx_bytes()
        resp = client.post(
            "/upload",
            files={
                "files": (
                    "smoke_test.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document",
                )
            },
        )
        assert resp.status_code == 200, f"POST /upload {resp.status_code}: {resp.text}"
        job_id = resp.json()["job_id"]
        assert job_id, "No job_id in upload response"

    run("B. Upload test DOCX (in-memory)", step_b)

    if job_id is None:
        print("\n  Cannot continue without a job. Aborting.")
        _summary()
        client.close()
        return

    # ------------------------------------------------------------------
    # C. Tier 2 anonymization via SSE stream
    # ------------------------------------------------------------------
    def step_c():
        with client.stream(
            "POST",
            f"/jobs/{job_id}/anonymize/stream",
            json={"tier": "names_patterns", "roster_id": roster_id},
        ) as resp:
            assert resp.status_code == 200, (
                f"POST /anonymize/stream {resp.status_code}"
            )
            completed = False
            for line in resp.iter_lines():
                if not line.startswith("data:"):
                    continue
                ev = json.loads(line[5:].strip())
                if ev.get("error"):
                    raise RuntimeError(f"Pipeline error: {ev['error']}")
                if ev.get("step") == "complete":
                    completed = True
                    break
        assert completed, "Stream closed without a 'complete' event"

    run("C. Tier 2 anonymization (SSE stream)", step_c)

    # ------------------------------------------------------------------
    # D. Review — verify all PII replaced
    # ------------------------------------------------------------------
    def step_d():
        resp = client.get(f"/jobs/{job_id}/review")
        assert resp.status_code == 200, f"GET /review {resp.status_code}"
        files = resp.json().get("files", [])
        assert files, "No files in review response"
        anon = files[0]["anonymized_text"]

        problems = []
        if "[PERSON_" not in anon:
            problems.append("No [PERSON_N] placeholder found")
        if f"{FIRST1} {LAST1}" in anon:
            problems.append(f"Original name '{FIRST1} {LAST1}' still present")
        if f"{FIRST2} {LAST2}" in anon:
            problems.append(f"Original name '{FIRST2} {LAST2}' still present")
        if "[EMAIL_" not in anon:
            problems.append("No [EMAIL_N] placeholder found")
        if DOCX_EMAIL in anon:
            problems.append(f"Original email still present")
        if "[PHONE_" not in anon:
            problems.append("No [PHONE_N] placeholder found")
        if DOCX_PHONE in anon:
            problems.append(f"Original phone still present")

        if problems:
            raise AssertionError(
                "; ".join(problems) + f"\n         anonymized_text: {anon[:300]}"
            )

    run("D. Review — names, email, phone all replaced", step_d)

    # ------------------------------------------------------------------
    # E. Export DOCX — verify placeholder in file
    # ------------------------------------------------------------------
    exported_bytes: bytes | None = None

    def step_e():
        nonlocal exported_bytes
        resp = client.get(f"/jobs/{job_id}/export")
        assert resp.status_code == 200, f"GET /export {resp.status_code}: {resp.text}"
        exported_bytes = resp.content
        assert len(exported_bytes) > 100, "Exported file is suspiciously small"

        full_text = _docx_full_text(exported_bytes)
        assert "[PERSON_" in full_text, (
            f"[PERSON_N] not in exported DOCX text. Got: {full_text[:300]}"
        )
        assert f"{FIRST1} {LAST1}" not in full_text, (
            "Original name still present in exported DOCX"
        )

    run("E. Export DOCX — valid file, placeholder present", step_e)

    # ------------------------------------------------------------------
    # F. Export mapping JSON
    # ------------------------------------------------------------------
    mapping_dict: dict | None = None

    def step_f():
        nonlocal mapping_dict
        resp = client.get(f"/jobs/{job_id}/export/mapping")
        assert resp.status_code == 200, f"GET /export/mapping {resp.status_code}"
        mappings = resp.json()
        assert mappings, "Mapping list is empty"
        mapping_dict = {m["placeholder"]: m["original"] for m in mappings}
        assert mapping_dict, "Could not build placeholder→original dict"
        # Sanity: at least one PERSON placeholder present
        person_keys = [k for k in mapping_dict if k.startswith("[PERSON_")]
        assert person_keys, f"No [PERSON_N] keys in mapping. Keys: {list(mapping_dict)}"

    run("F. Export mapping JSON", step_f)

    # ------------------------------------------------------------------
    # G. Re-identify — verify original names restored in DOCX
    # ------------------------------------------------------------------
    def step_g():
        assert exported_bytes is not None, "No exported bytes (step E failed)"
        assert mapping_dict is not None, "No mapping (step F failed)"

        resp = client.post(
            "/reidentify",
            json={"job_id": job_id, "mapping": mapping_dict},
        )
        assert resp.status_code == 200, (
            f"POST /reidentify {resp.status_code}: {resp.text}"
        )

        restored_text = _docx_full_text(resp.content)
        rt_lower = restored_text.lower()
        assert LAST1.lower() in rt_lower or FIRST1.lower() in rt_lower, (
            f"Original name not restored. Got: {restored_text[:300]}"
        )
        person_keys = [k for k in mapping_dict if k.startswith("[PERSON_")]
        for key in person_keys:
            assert key not in restored_text, (
                f"Placeholder {key!r} still in re-identified DOCX"
            )

    run("G. Re-identify — original names restored in DOCX", step_g)

    # ------------------------------------------------------------------
    # I. Upload test PDF
    # ------------------------------------------------------------------
    pdf_job_id: str | None = None

    def step_i():
        nonlocal pdf_job_id
        pdf_bytes = _make_pdf_bytes()
        resp = client.post(
            "/upload",
            files={"files": ("smoke_test.pdf", pdf_bytes, "application/pdf")},
        )
        assert resp.status_code == 200, f"POST /upload (PDF) {resp.status_code}: {resp.text}"
        pdf_job_id = resp.json()["job_id"]
        assert pdf_job_id, "No job_id in PDF upload response"

    run("I. Upload test PDF (in-memory)", step_i)

    if pdf_job_id is None:
        print("\n  Cannot continue PDF steps without a job. Skipping J-L.")
    else:
        # ------------------------------------------------------------------
        # J. Tier 2 anonymization of PDF
        # ------------------------------------------------------------------
        def step_j():
            with client.stream(
                "POST",
                f"/jobs/{pdf_job_id}/anonymize/stream",
                json={"tier": "names_patterns", "roster_id": roster_id},
            ) as resp:
                assert resp.status_code == 200, (
                    f"POST /anonymize/stream (PDF) {resp.status_code}"
                )
                completed = False
                for line in resp.iter_lines():
                    if not line.startswith("data:"):
                        continue
                    ev = json.loads(line[5:].strip())
                    if ev.get("error"):
                        raise RuntimeError(f"Pipeline error: {ev['error']}")
                    if ev.get("step") == "complete":
                        completed = True
                        break
            assert completed, "PDF stream closed without a 'complete' event"

        run("J. Tier 2 anonymization of PDF (SSE stream)", step_j)

        # ------------------------------------------------------------------
        # K. Review PDF — verify PII replaced
        # ------------------------------------------------------------------
        def step_k():
            resp = client.get(f"/jobs/{pdf_job_id}/review")
            assert resp.status_code == 200, f"GET /review (PDF) {resp.status_code}"
            files = resp.json().get("files", [])
            assert files, "No files in PDF review response"
            anon = files[0]["anonymized_text"]

            problems = []
            if "[PERSON_" not in anon:
                problems.append("No [PERSON_N] placeholder found")
            if f"{FIRST1} {LAST1}" in anon:
                problems.append(f"Original name '{FIRST1} {LAST1}' still present")
            if "[EMAIL_" not in anon:
                problems.append("No [EMAIL_N] placeholder found")
            if PDF_EMAIL in anon:
                problems.append("Original PDF email still present")
            if "[PHONE_" not in anon:
                problems.append("No [PHONE_N] placeholder found")
            if PDF_PHONE in anon:
                problems.append("Original PDF phone still present")

            if problems:
                raise AssertionError(
                    "; ".join(problems) + f"\n         anonymized_text: {anon[:300]}"
                )

        run("K. Review PDF — names, email, phone all replaced", step_k)

        # ------------------------------------------------------------------
        # L. Export PDF — valid file with placeholders
        # ------------------------------------------------------------------
        def step_l():
            resp = client.get(f"/jobs/{pdf_job_id}/export")
            assert resp.status_code == 200, f"GET /export (PDF) {resp.status_code}: {resp.text}"
            pdf_data = resp.content
            assert len(pdf_data) > 100, "Exported PDF is suspiciously small"

            full_text = _pdf_full_text(pdf_data)
            assert "[PERSON_" in full_text, (
                f"[PERSON_N] not in exported PDF text. Got: {full_text[:300]}"
            )
            assert f"{FIRST1} {LAST1}" not in full_text, (
                "Original name still present in exported PDF"
            )

        run("L. Export PDF — valid file, placeholder present", step_l)

        # ------------------------------------------------------------------
        # N. Re-identify PDF — verify originals restored
        # ------------------------------------------------------------------
        def step_n():
            # Fetch the mapping for the PDF job
            resp = client.get(f"/jobs/{pdf_job_id}/export/mapping")
            assert resp.status_code == 200, f"GET /export/mapping (PDF) {resp.status_code}"
            mappings = resp.json()
            assert mappings, "PDF mapping list is empty"
            pdf_mapping = {m["placeholder"]: m["original"] for m in mappings}

            resp2 = client.post(
                "/reidentify",
                json={"job_id": pdf_job_id, "mapping": pdf_mapping},
            )
            assert resp2.status_code == 200, (
                f"POST /reidentify (PDF) {resp2.status_code}: {resp2.text}"
            )

            restored_text = _pdf_full_text(resp2.content)
            rt_lower = restored_text.lower()
            assert LAST1.lower() in rt_lower or FIRST1.lower() in rt_lower, (
                f"Original name not restored in PDF. Got: {restored_text[:300]}"
            )
            person_keys = [k for k in pdf_mapping if k.startswith("[PERSON_")]
            for key in person_keys:
                assert key not in restored_text, (
                    f"Placeholder {key!r} still in re-identified PDF"
                )

        run("N. Re-identify PDF — original names restored", step_n)

    # ------------------------------------------------------------------
    # O. Discover — quick scan of a DOCX returns known PII findings
    # ------------------------------------------------------------------
    discover_findings: list | None = None

    def step_o():
        nonlocal discover_findings
        docx_bytes = _make_docx_bytes()
        resp = client.post(
            "/discover",
            files={
                "file": (
                    "discover_test.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document",
                )
            },
            data={"method": "quick"},
        )
        assert resp.status_code == 200, f"POST /discover {resp.status_code}: {resp.text}"
        data = resp.json()
        assert "findings" in data, "No 'findings' key in /discover response"
        assert isinstance(data["findings"], list), "findings is not a list"
        assert "warnings" in data, "No 'warnings' key in /discover response"
        assert "job_id" not in data, "/discover must not return a job_id (must be stateless)"
        discover_findings = data["findings"]

    run("O. Discover — quick scan of DOCX (stateless)", step_o)

    # ------------------------------------------------------------------
    # P. Discover — verify findings contain expected PII + CSV format
    # ------------------------------------------------------------------
    def step_p():
        assert discover_findings is not None, "No findings from step O"
        pii_types = {f["pii_type"] for f in discover_findings}
        texts = {f["text"] for f in discover_findings}

        assert "EMAIL" in pii_types, (
            f"EMAIL not found in discovered types. Found: {pii_types}"
        )
        assert DOCX_EMAIL in texts, (
            f"Email {DOCX_EMAIL!r} not found in findings. Texts: {list(texts)[:10]}"
        )
        assert "PHONE" in pii_types, (
            f"PHONE not found in discovered types. Found: {pii_types}"
        )
        assert any(DOCX_PHONE in t for t in texts), (
            f"Phone {DOCX_PHONE!r} not found in findings. Texts: {list(texts)[:10]}"
        )

        # Verify each finding has the required schema fields
        for f in discover_findings:
            for field in ("text", "pii_type", "confidence", "source"):
                assert field in f, f"Finding missing field {field!r}: {f}"

        # Verify CSV column assignment logic for email and phone findings
        import csv, io as _io
        rows_out = [["first_name", "last_name", "email", "also_remove"]]
        for f in discover_findings:
            row = ["", "", "", ""]
            if f["pii_type"] == "EMAIL":
                row[2] = f["text"]
            elif f["pii_type"] == "PERSON":
                parts = f["text"].split(" ", 1)
                row[0] = parts[0]
                row[1] = parts[1] if len(parts) > 1 else ""
            else:
                row[3] = f["text"]
            rows_out.append(row)

        buf = _io.StringIO()
        csv.writer(buf, lineterminator="\r\n").writerows(rows_out)
        csv_text = buf.getvalue()

        assert DOCX_EMAIL in csv_text, (
            f"Email not in generated CSV.\nCSV:\n{csv_text[:400]}"
        )
        assert "first_name,last_name,email,also_remove" in csv_text, (
            "CSV header row missing or malformed"
        )
        # Email must appear in the third column (email), not also_remove
        for line in csv_text.splitlines():
            if DOCX_EMAIL in line:
                cols = line.split(",")
                assert cols[2] == DOCX_EMAIL, (
                    f"Email should be in column index 2 (email), got: {cols}"
                )

    run("P. Discover — findings contain email + phone; CSV format correct", step_p)

    # ------------------------------------------------------------------
    # M. Shutdown
    # ------------------------------------------------------------------
    if _SHUTDOWN:
        def step_m():
            resp = client.post("/shutdown")
            assert resp.status_code == 200, f"POST /shutdown {resp.status_code}"

        run("M. POST /shutdown", step_m)

    client.close()
    _summary()


def _summary() -> None:
    total = len(_results)
    passed = sum(1 for _, ok, _ in _results if ok)
    print()
    print(f"  Result: {passed}/{total} passed")
    if passed < total:
        sys.exit(1)


if __name__ == "__main__":
    main()
